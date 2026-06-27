#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx — تبدیل قالب‌های Markdown مستندسازی فرایند به Word حرفه‌ای (راست‌به‌چپ فارسی).

پشتیبانی: عنوان‌ها (#..######)، جدول‌ها (| ... |)، لیست‌ها (- / *)،
نقل‌قول راهنما (>)، خط جداکننده (---)، متن ساده، و رد کردن کامنت‌های HTML.

استفاده:
    python md2docx.py <input.md> [output.docx]
    python md2docx.py --all            # تبدیل همه‌ی قالب‌های ProcessDocKit/templates → ProcessDocKit/word
"""
import sys
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "B Nazanin"          # فونت اصلی فارسی؛ در نبود آن Word به فونت پیش‌فرض برمی‌گردد
FONT_FALLBACK = "Tahoma"
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
TABLE_HEADER_BG = "1F3A5F"
TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)


def set_rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def style_run(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)
    # اندازه‌ی complex-script (برای فارسی)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(szcs)
    rtl = OxmlElement("w:rtl")
    rpr.append(rtl)


def add_paragraph(doc, text, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    for seg, is_bold in parse_bold(text):
        run = p.add_run(seg)
        style_run(run, size=size, bold=bold or is_bold, color=color)
    return p


def parse_bold(text):
    """تقسیم متن بر اساس **bold** به قطعات (متن، آیا bold)."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    out = []
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            out.append((part[2:-2], True))
        elif part:
            out.append((part, False))
    if not out:
        out = [("", False)]
    return out


def clean_inline(text):
    text = re.sub(r"`([^`]*)`", r"\1", text)            # حذف بک‌تیک
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # لینک → فقط متن
    return text.strip()


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)


def fill_cell(cell, text, bold=False, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    color = TABLE_HEADER_FG if header else None
    for seg, is_bold in parse_bold(text):
        run = p.add_run(seg)
        style_run(run, size=10, bold=bold or is_bold or header, color=color)
    if header:
        shade_cell(cell, TABLE_HEADER_BG)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells) and len(cells) > 0


def split_row(line):
    return [clean_inline(c.strip()) for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)
    for i, h in enumerate(header):
        fill_cell(table.rows[0].cells[i], h, header=True)
    for r in body:
        cells = table.add_row().cells
        for i in range(ncol):
            val = r[i] if i < len(r) else ""
            fill_cell(cells[i], val)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        text = f.read()

    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)  # حذف کامنت‌های HTML
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.right_margin = Cm(2)
    section.left_margin = Cm(2)

    i = 0
    n = len(lines)
    quote_buffer = []

    def flush_quote():
        if quote_buffer:
            joined = " ".join(quote_buffer)
            p = add_paragraph(doc, joined, size=10, color=RGBColor(0x55, 0x55, 0x55))
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            quote_buffer.clear()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_quote()
            i += 1
            continue

        # نقل‌قول (راهنما)
        if stripped.startswith(">"):
            quote_buffer.append(clean_inline(stripped.lstrip(">").strip()))
            i += 1
            continue
        else:
            flush_quote()

        # خط جداکننده
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # عنوان‌ها
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            content = clean_inline(m.group(2))
            sizes = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11, 6: 11}
            p = add_paragraph(doc, content, size=sizes.get(level, 11),
                              bold=True, color=HEADING_COLOR, space_after=8)
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            i += 1
            continue

        # جدول
        if is_table_row(line):
            block = []
            while i < n and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            rows = [split_row(b) for b in block if not is_separator_row(b)]
            if rows:
                add_table(doc, rows)
            continue

        # لیست
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            content = clean_inline(m.group(1))
            p = doc.add_paragraph(style="List Bullet")
            set_rtl_paragraph(p)
            for seg, is_bold in parse_bold(content):
                run = p.add_run(seg)
                style_run(run, size=11, bold=is_bold)
            i += 1
            continue

        # متن ساده
        add_paragraph(doc, clean_inline(stripped))
        i += 1

    flush_quote()
    doc.save(docx_path)
    return docx_path


def main():
    args = sys.argv[1:]
    here = os.path.dirname(os.path.abspath(__file__))
    kit_dir = os.path.normpath(os.path.join(here, ".."))
    templates_dir = os.path.join(kit_dir, "templates")

    if not args:
        print(__doc__)
        return

    if args[0] == "--all":
        out_dir = os.path.join(kit_dir, "word")
        os.makedirs(out_dir, exist_ok=True)
        for fn in sorted(os.listdir(templates_dir)):
            if fn.endswith(".md"):
                src = os.path.join(templates_dir, fn)
                dst = os.path.join(out_dir, fn[:-3] + ".docx")
                convert(src, dst)
                print("OK:", dst)
        return

    src = args[0]
    dst = args[1] if len(args) > 1 else os.path.splitext(src)[0] + ".docx"
    convert(src, dst)
    print("OK:", dst)


if __name__ == "__main__":
    main()
